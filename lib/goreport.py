#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""This is the GoReport class. GoReport handles everything from connecting to the target GoPhish
server to pulling campaign information and reporting the results. GoReport only requires a few
variables when a new GoReport object is created. There is no need to call specific functions
or do anything beyond creating the object, like so:

gophish = goreport.GoReport(format, config, verbose)
"""

# Try to import gophish
try:
    from gophish import Gophish
except:
    print("[!] Could not import the GoPhish library! Make sure it is installed.\n\
Run: `python3 -m pip intall gophish`\n\
Test it by running `python3` and then, in the \
Python prompt, typing `from gophish import Gophish`.")
    exit()

# Imports for statistics, e.g. browser and operating systems
from user_agents import parse
from collections import Counter, OrderedDict
# Import the MaxMind's GeoLite for IP address GeoIP look-ups
from geolite2 import geolite2
# Imports for writing the Word.doc report
import os.path
from docx import *
from docx.shared import *
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.shared import OxmlElement, qn
# Basic imports
import sys
import csv
import configparser
import time
# Imports for web requests, e.g. Google Maps API for location data
# Disables the insecure HTTPS warning for the self-signed GoPhish certs
import requests
from requests.packages.urllib3.exceptions import InsecureRequestWarning
requests.packages.urllib3.disable_warnings(InsecureRequestWarning)



class GoReport(object):
    """This class uses the GoPhish library to create a new GoPhish API connection
    and queries GoPhish for information and results related to the specified
    campaign ID(s).
    """
    # Name of the config file -- default is gophish.config
    goreport_config_file = "gophish.config"
    verbose = False

    # Variables for holding GoPhish models
    campaign = None
    results = None
    timeline = None

    # Variables for holding campaign information
    cam_id = None
    cam_name = None
    cam_status = None
    created_date = None
    launch_date = None
    completed_date = None
    cam_url = None
    cam_redirect_url = None
    cam_from_address = None
    cam_subject_line = None
    cam_template_name = None
    cam_capturing_passwords = None
    cam_capturing_credentials = None
    cam_page_name = None
    cam_smtp_host = None

    # Variables and lists for tracking event numbers
    total_targets = 0
    total_sent = 0
    total_opened = 0
    total_clicked = 0
    total_submitted = 0
    total_unique_opened = 0
    total_unique_clicked = 0
    total_unique_submitted = 0
    targets_opened = []
    targets_clicked = []
    targets_submitted = []

    # Lists and dicts for holding prepared report data
    campaign_results_summary = []

    # Lists for holding totals for statistics
    browsers = []
    operating_systems = []
    locations = []
    ip_addresses = []

    # Output options
    report_format = None
    output_csv_report = None
    output_word_report = None

    def __init__(self, report_format, config_file, verbose):
        """Initiate the connection to the GoPhish server with the provided host,
        port, and API key.
        """
        try:
            # Check if an alternate config file was provided
            if config_file is not None:
                self.goreport_config_file = config_file
            # Open the config file to make sure it exists and is readable
            config = configparser.ConfigParser()
            config.read(self.goreport_config_file)
        except Exception:
            print("[!] Could not open the /gophish.config config file \
-- make sure it exists and is readable.")
            print("L.. Details: {}".format(e))
            sys.exit()

        try:
            # Read in the values from the config file
            GP_HOST = self.config_section_map(config, "GoPhish")['gp_host']
            API_KEY = self.config_section_map(config, "GoPhish")['api_key']
        except Exception as e:
            print("[!] There was a problem reading values from the gophish.config file!")
            print("L.. Details: {}".format(e))
            sys.exit()

        # Set command line options for the GoReport object
        self.verbose = verbose
        self.report_format = report_format
        # Create the MaxMind GeoIP reader for MaxMind lookups
        self.geoip_reader = geolite2.reader()
        # Connect to the GoPhish API
        # NOTE: This step succeeds even with a bad API key, so
        # the true test is fetching an ID in run()
        print("[+] Connecting to GoPhish at {}".format(GP_HOST))
        print("L.. The API Authorization endpoint is: \
{}/api/campaigns/?api_key={}".format(GP_HOST, API_KEY))
        self.api = Gophish(API_KEY, host=GP_HOST, verify=False)

    def run(self, id_list, combine_reports, set_complete_status):
        """Run everything to process the target campaign."""
        # Output some feedback for user options
        if combine_reports:
            print("[+] Campaign results will be combined into a single report.")
        if set_complete_status:
            print('[+] Campaign statuses will be set to "Complete" after processing the results.')

        try:
            # Create the list of campaign IDs from --id
            temp_id = []
            # Handle a mixed set of ranges and comma-separated IDs
            if "-" and "," in id_list:
                temp = id_list.split(",")
                for x in temp:
                    if "-" in x:
                        lower = x.split("-")[0]
                        upper = x.split("-")[1]
                        for y in range(int(lower), int(upper) + 1):
                            temp_id.append(str(y))
                    else:
                        temp_id.append(x)
            # Process IDs provided as one or more ranges
            elif "-" in id_list:
                lower = id_list.split("-")[0]
                upper = id_list.split("-")[1]
                for y in range(int(lower), int(upper) + 1):
                    temp_id.append(str(y))
            # Process single or only comma-separated IDs
            else:
                temp_id = id_list.split(",")
            id_list = temp_id
        except Exception as e:
            print("[!] Could not interpret your provided campaign IDs. \
Ensure the IDs are provided as comma-separated integers or interger ranges, e.g. 5,50-55,71.")
            print("L.. Details: {}".format(e))
            sys.exit()

        # Begin proessing the campaign IDs by removing any duplicates
        try:
            # Get length of user-provided list
            initial_len = len(id_list)
            # Remove duplicate IDs and sort IDs as integers
            id_list = sorted(set(id_list), key=int)
            # Get length of unique, sorted list
            unique_len = len(id_list)
        except Exception as e:
            temp = []
            for id in id_list:
                try:
                    int(id)
                except:
                    temp.append(id)
            print("[!] There are {} invalid campaign ID(s), i.e. not an integer. \
Please correct the following ID(s).".format(len(temp)))
            print("L.. Offending IDs: {}".format(",".join(temp)))
            sys.exit()
        print("[+] A total of {} campaign IDs have been provided for \
processing.".format(initial_len))
        # If the lengths are different, then GoReport removed one or more dupes
        if initial_len != unique_len:
            dupes = initial_len - unique_len
            print("L.. GoReport found {} duplicate campaign IDs, so those \
have been trimmed to avoid bad results with --combine and wasted time.".format(dupes))
        # Provide  list of all IDs that will be processed
        print("[+] GoReport will process IDs {}".format(",".join(id_list)))

        # If --combine is used with just one ID it can break reporting, so we catch that here
        if len(id_list) == 1 and combine_reports:
            print("[+] You provided just one campaign ID and enabled combining \
reports, so GoReport is going to ignore --combine.")
            combine_reports = False

        print()

        # Go through each campaign ID and get the results
        campaign_counter = 1
        for CAM_ID in id_list:
            print("[+] Now fetching results for Campaign ID {} \
({}/{}).".format(CAM_ID, campaign_counter, len(id_list)))
            try:
                # Request the details for the provided campaign ID
                self.campaign = self.api.campaigns.get(campaign_id=CAM_ID)
            except Exception as e:
                print("[!] There was a problem fetching this campaign ID's details. \
Make sure your URL and API key are correct. Check HTTP vs HTTPS!".format(CAM_ID))
                print("L.. Details: {}".format(e))
                sys.exit()

            try:
                try:
                    # Check to see if a success message was returned with a message
                    # Possible reasons: campaign ID doesn't exist or problem with host/API key
                    if self.campaign.success is False:
                        print("[!] Failed to get results for campaign ID {}".format(CAM_ID))
                        print("L.. Details: {}".format(self.campaign.message))

                        # We can't let an error with an ID stop reporting, so
                        # check if this was the last ID for a combiend report
                        if CAM_ID == id_list[-1] and combine_reports:
                            self.generate_report()
                # If self.campaign.success does not exist then we were successful
                except:
                    print("[+] Success!")
                    # Collect campaign details and process data
                    self.collect_all_campaign_info(combine_reports)
                    self.process_timeline_events(combine_reports)
                    self.process_results(combine_reports)

                    # Check if this is the last campaign ID in the list
                    # If this is the last ID and combined reports is on, generate the report
                    if CAM_ID == id_list[-1] and combine_reports:
                        self.generate_report()
                    # Otherwise, if we are not combining reports, generate the reports
                    elif combine_reports is False:
                        self.generate_report()
                    campaign_counter += 1

                # If the --complete flag was set, now set campaign status to Complete
                if set_complete_status:
                    print("[+] Setting campaign ID {}'s status to Complete.".format(CAM_ID))
                    try:
                        set_complete = self.api.campaigns.complete(CAM_ID)
                        try:
                            if set_complete.success is False:
                                print("[!] Failed to set campaign status for ID {}.".format(CAM_ID))
                                print("L.. Details: {}".format(set_complete.message))
                        # If set_complete.success does not exist then we were successful
                        except:
                            pass
                    except Exception as e:
                        print("[!] Failed to set campaign status for ID {}.".format(CAM_ID))
                        print("L.. Details: {}".format(e))
            except Exception as e:
                print("[!] There was a problem processing campaign ID {}!".format(CAM_ID))
                print("L.. Details: {}".format(e))
                sys.exit()

    def generate_report(self):
        """This function determines which report type to generate and then calls
        the appropriate reporting functions.
        """
        if self.report_format == "csv":
            print("[+] Building the report -- you selected a csv report.")
            self.output_csv_report = self._build_output_csv_file_name()
            self.write_csv_report()
        elif self.report_format == "word":
            print("[+] Building the report -- you selected a Word/docx report.")
            print("[+] Looking for the template.docx to be used for the Word report.")
            if os.path.isfile("template.docx"):
                print("[+] Template was found -- proceeding with report generation...")
                print("L.. This may take a while if you provided a lot of \
IDs for a combined report or have a lot of targets.")
                self.output_word_report = self._build_output_word_file_name()
                self.write_word_report()
            else:
                print("[!] Could not find the template document! Make sure \
'template.docx' is in the GoReport directory.")
                sys.exit()
        elif self.report_format == "quick":
            print("[+] Quick report stats:")
            self.get_quick_stats()

    def get_quick_stats(self):
        """Present quick stats for the campaign. Just basic numbers and some details."""
        print("")
        print(self.cam_name)
        print("Status:\t\t{}".format(self.cam_status))
        print("Created:\t{} on {}".format(self.created_date.split("T")[1].split(".")[0],
                                          self.created_date.split("T")[0]))
        print("Started:\t{} on {}".format(self.launch_date.split("T")[1].split(".")[0],
                                          self.launch_date.split("T")[0]))
        if self.cam_status == "Completed":
            print("Completed:\t{} on {}".format(self.completed_date.split("T")[1].split(".")[0],
                                                self.completed_date.split("T")[0]))
        print("")
        print("Total Targets:\t{}".format(self.total_targets))
        print("Emails Sent:\t{}".format(self.total_sent))
        print("IPs Seen:\t{}".format(len(self.ip_addresses)))
        print("")
        print("Total Opened Events:\t\t{}".format(self.total_opened))
        print("Total Click Events:\t\t{}".format(self.total_clicked))
        print("Total Submitted Data Events:\t{}".format(self.total_submitted))
        print("")
        print("Individuals Who Opened:\t\t{}".format(self.total_unique_opened))
        print("Individuals Who Clicked:\t{}".format(self.total_unique_clicked))
        print("Individuals Who Entered data:\t{}".format(self.total_unique_submitted))

    def _build_output_csv_file_name(self):
        """A helper function to create the output report name."""
        csv_report = "GoPhish Results for Campaign - {}.csv".format(self.cam_name)
        return csv_report

    def _build_output_word_file_name(self):
        """A helper function to create the output report name."""
        word_report = "GoPhish Results for Campaign - {}.docx".format(self.cam_name)
        return word_report

    def set_word_column_width(self, column, width):
        """Custom function for quickly and easily setting the width of a table's
        column in the Word docx output.

        An option missing from the basic Python docx library.
        """
        for cell in column.cells:
            cell.width = width

    def get_basic_campaign_info(self):
        """"Helper function to collect a campaign's basic details. This includes
        campaign name, status, template, and other details that are not the
        campaign's results.

        This keeps these calls in one place for tidiness and easier management.
        """
        self.cam_name = self.campaign.name
        self.cam_status = self.campaign.status
        self.created_date = self.campaign.created_date
        self.launch_date = self.campaign.launch_date
        self.completed_date = self.campaign.completed_date
        self.cam_url = self.campaign.url

        # Collect SMTP information
        self.smtp = self.campaign.smtp
        self.cam_from_address = self.smtp.from_address
        self.cam_smtp_host = self.smtp.host

        # Collect the template information
        self.template = self.campaign.template
        self.cam_subject_line = self.template.subject
        self.cam_template_name = self.template.name
        self.cam_template_attachments = self.template.attachments
        if self.cam_template_attachments == []:
            self.cam_template_attachments = "None Used"

        # Collect the landing page information
        self.page = self.campaign.page
        self.cam_page_name = self.page.name
        self.cam_redirect_url = self.page.redirect_url
        if self.cam_redirect_url == "":
            self.cam_redirect_url = "Not Used"
        self.cam_capturing_passwords = self.page.capture_passwords
        self.cam_capturing_credentials = self.page.capture_credentials

    def collect_all_campaign_info(self, combine_reports):
        """Collect the campaign's details set values for each of the declared variables."""
        # Collect the basic campaign details
        # Plus a quick and dirty check to see if the campaign ID is valid
        try:
            self.cam_id = self.campaign.id
        except:
            print("[!] Looks like campaign ID {} does not exist! \
Skipping it...".format(self.cam_id))

        if combine_reports and self.cam_name is None:
            print("[+] Reports will be combined -- setting name, dates, and \
URL based on campaign ID {}.".format(self.cam_id))
            self.get_basic_campaign_info()
        elif combine_reports is False:
            self.get_basic_campaign_info()

        # Collect the results and timeline lists
        if self.results is None:
            self.results = self.campaign.results
            self.timeline = self.campaign.timeline
        elif combine_reports:
            self.results += self.campaign.results
            self.timeline += self.campaign.timeline
        else:
            self.results = self.campaign.results
            self.timeline = self.campaign.timeline

    def process_results(self, combine_reports):
        """Process the results model to collect basic data, like total targets
        and who opened, clicked, or submitted data. This should be run after
        the process_timeline_events() function which creates the targets_* lists.

        The results model can provide:
        first_name, last_name, email, position, and IP address
        """
        # Total length of results gives us the total number of targets
        if combine_reports and self.total_targets is None:
            self.total_targets = len(self.campaign.results)
        elif combine_reports:
            self.total_targets += len(self.campaign.results)
        else:
            # Reports will not be combined, so reset tracking between reports
            self.total_targets = len(self.campaign.results)
            self.ip_addresses = []
            self.campaign_results_summary = []

        # Go through all results and extract data for statistics
        for target in self.campaign.results:
            temp_dict = {}

            # Log the IP address for additional statistics later
            if not target.ip == "":
                self.ip_addresses.append(target.ip)
            # Add all of the recipient's details and results to the temp dictionary
            temp_dict["email"] = target.email
            temp_dict["fname"] = target.first_name
            temp_dict["lname"] = target.last_name
            temp_dict["ip_address"] = target.ip
            # Chck if this target was recorded as viewing the email (tracking image)
            if target.email in self.targets_opened:
                temp_dict["opened"] = "Y"
                self.total_unique_opened += 1
            else:
                temp_dict["opened"] = "-"
            # Check if this target clicked the link
            if target.email in self.targets_clicked:
                temp_dict["clicked"] = "Y"
                self.total_unique_clicked += 1
            else:
                temp_dict["clicked"] = "-"
            # Check if this target submitted data
            if target.email in self.targets_submitted:
                temp_dict["submitted"] = "Y"
                self.total_unique_submitted += 1
            else:
                temp_dict["submitted"] = "-"
            # Append the temp dictionary to the event summary list
            self.campaign_results_summary.append(temp_dict)

    def process_timeline_events(self, combine_reports):
        """Process the timeline model to collect basic data, like total clicks,
        and get detailed event data for recipients.

        The timeline model contains all events that occured during the campaign.
        """
        # Create counters for enumeration
        sent_counter = 0
        opened_counter = 0
        click_counter = 0
        submitted_counter = 0
        # Run through all events and count each of the four basic events
        for event in self.campaign.timeline:
            if event.message == "Email Sent":
                sent_counter += 1
            elif event.message == "Email Opened":
                opened_counter += 1
                self.targets_opened.append(event.email)
            elif event.message == "Clicked Link":
                click_counter += 1
                self.targets_clicked.append(event.email)
            elif event.message == "Submitted Data":
                submitted_counter += 1
                self.targets_submitted.append(event.email)
        # Assign the counter values to our tracking lists
        if combine_reports:
            # Append, +=, totals if combining reports
            self.total_sent += sent_counter
            self.total_opened += opened_counter
            self.total_clicked += click_counter
            self.total_submitted += submitted_counter
        else:
            # Set tracking variables to current counter values for non-combined reports
            self.total_sent = sent_counter
            self.total_opened = opened_counter
            self.total_clicked = click_counter
            self.total_submitted = submitted_counter

    def lookup_ip(self, ip):
        """Check the GeoLite database for a location for the provided IP address.

        This returns a large dict with more data than is probably needed for
        a report. This gets continent, country, registered_country, and location.
        Also, this dict includes multiple languages.

        You may wonder why get_google_location_data() is needed if this provides
        a lot of data from MaxMind. Unfortunately, the MaxMind database will not
        always have the data needed most for the report (city, state, country).
        It may only have the continent name. Luckily, it seems to always have coordinates
        that can be compared to GoPhish's coordinates and passed to get_google_location_data().
        """
        match = self.geoip_reader.get(ip)
        if match is not None:
            return match
        else:
            # return "No match"
            return None

    def get_google_location_data(self, lat, lon):
        """Use Google's Maps API to collect GeoIP info for the provided latitude
        and longitude.

        Google returns a bunch of JSON with a variety of location data.
        This function sticks to the first set of "address_components" for the
        country, locality (city), and administrative_level_1 (state).

        Ex: http://maps.googleapis.com/maps/api/geocode/json?latlng=35,-93&sensor=false
        """
        url = "http://maps.googleapis.com/maps/api/geocode/json?latlng={},\
{}&sensor=false".format(lat, lon)
        v = requests.get(url)
        j = v.json()
        try:
            # Get the first set of 'address_components' from the JSON results
            components = j['results'][0]['address_components']
            country = town = None
            for c in components:
                if "country" in c['types']:
                    country = c['long_name']
                if "locality" in c['types']:
                    town = c['long_name']
                if "administrative_area_level_1" in c['types']:
                    state = c['long_name']
            return "{} {} {}".format(town, state, country)
        except:
            # return "None"
            return None

    def compare_ip_addresses(self, target_ip, browser_ip, verbose):
        """Compare the IP addresses of the target to that of an event. The goal:
        Looking for a mismatch that might identify some sort of interesting event.
        This might indicate an email was forwarded, a VPN was switched on/off, or
        maybe the target is at home.
        """
        if target_ip == browser_ip:
            return target_ip
        else:
            # We have an IP mismatch -- hard to tell why this might be.
            if verbose:
                print("[*] Event: This target's ({}) URL was clicked from a \
browser at {}.".format(target_ip, browser_ip))
            # This is an IP address not included in the results model, so we add it to our list here
            self.ip_addresses.append(browser_ip)
            return browser_ip

    def compare_ip_coordinates(self, target_latitude, target_longitude, mmdb_latitude,
                               mmdb_longitude, ip_address, verbose):
        """Compare the IP address cooridnates reported by MaxMind and GoPhish.
        If they do not match, some additional -- manual -- investigation should
        be done for any client-facing deliverables.
        """
        if target_latitude == mmdb_latitude and target_longitude == mmdb_longitude:
            # Coordinates match what GoPhish recorded, so query Google Maps for details
            coordinates_location = self.get_google_location_data(target_latitude, target_longitude)
            if not coordinates_location is None:
                self.locations.append(coordinates_location)
                return coordinates_location
            else:
                return "Google timeout"
        else:
            # MaxMind and GoPhish have different coordinates, so this is a tough spot
            # Both locations can be recorded for investigation, but what to do
			# for location statistics? It was decided both would be recorded as
			# one location with an asterisk, flagged for investigation.
            if verbose:
                print("[*] Warning: Location coordinates mis-match between MaxMind and GoPhish \
for {}. Check report for location with * to investigate and pick the right \
location.".format(ip_address))
            coordinates_location = self.get_google_location_data(target_latitude, target_longitude)
            # Sleep between checks to avoid timeouts
            time.sleep(2)
            alt_coordinates_location = self.get_google_location_data(mmdb_latitude, mmdb_longitude)
            if not alt_coordinates_location is None and not coordinates_location is None:
                coordinates_location += " ALTERNATE:{}".format(alt_coordinates_location)
            elif not coordinates_location is None and alt_coordinates_location is None:
                coordinates_location += "     ALTERNATE: MaxMind returned No Results"
            elif coordinates_location is None and not alt_coordinates_location is None:
                coordinates_location = alt_coordinates_location

            try:
                self.locations.append(coordinates_location + " *")
                return "{}".format(coordinates_location + " *")
            except Exception:
                print("[!] There was a problem getting location for IP: {}".format(ip_address))
                coordinates_location = "Error"
                alt_coordinates_location = "Error"
                self.locations.append(coordinates_location + " *")
                return "{}".format(coordinates_location + " *")

    def write_csv_report(self):
        """Assemble and output the csv file report.

        Throughout this function, results are assembled by adding commas and then
        adding to a results string, i.e. 'result_A' and then 'result_A' += ',result_B'.
        This is so the result can be written to the csv file and have the different
        pieces end up in the correct columns.
        """
        with open(self.output_csv_report, 'w') as csvfile:
            # Create the csv writer
            writer = csv.writer(csvfile, dialect='excel', delimiter=',', quotechar=" ",
                                quoting=csv.QUOTE_MINIMAL)

            # Write a campaign summary at the top of the report
            writer.writerow(["Campaign Results For:", "{}".format(self.cam_name)])
            writer.writerow(["Status", "{}".format(self.cam_status)])
            writer.writerow(["Created", "{}".format(self.created_date)])
            writer.writerow(["Started", "{}".format(self.launch_date)])
            # If the campaign has been completed, we will record that, too
            if self.cam_status == "Completed":
                writer.writerow(["Completed", "{}".format(self.completed_date)])
            # Write the campaign details -- email details and template settings
            writer.writerow("")
            writer.writerow(["Campaign Details"])
            writer.writerow(["From", "{}".format(self.cam_from_address)])
            writer.writerow(["Subject", "{}".format(self.cam_subject_line)])
            writer.writerow(["Phish URL", "{}".format(self.cam_url)])
            writer.writerow(["Redirect URL", "{}".format(self.cam_redirect_url)])
            writer.writerow(["Attachment(s)", "{}".format(self.cam_template_attachments)])
            writer.writerow(["Captured Credentials", "{}".format(self.cam_capturing_credentials)])
            writer.writerow(["Stored Passwords", "{}".format(self.cam_capturing_passwords)])
            # Write a high level summary for stats
            writer.writerow("")
            writer.writerow(["High Level Results"])
            writer.writerow(["Total Targets", "{}".format(self.total_targets)])
            writer.writerow("")
            writer.writerow(["The following totals indicates how many events of each type \
                            GoPhish recorded:"])
            writer.writerow(["Total Opened Events", "{}".format(self.total_opened)])
            writer.writerow(["Total Clicked Events", "{}".format(self.total_clicked)])
            writer.writerow(["Total Submitted Data Events", "{}".format(self.total_submitted)])
            writer.writerow("")
            writer.writerow(["The following totals indicates how many targets participated in \
                            each event type:"])
            writer.writerow(["Individuals Who Opened", "{}".format(self.total_unique_opened)])
            writer.writerow(["Individuals Who Clicked", "{}".format(self.total_unique_clicked)])
            writer.writerow(["Individuals Who Submitted Data", "{}".
                             format(self.total_unique_submitted)])

            print("[+] Finished writing high level summary...")
            # End of the campaign summary and beginning of the event summary
            writer.writerow("")
            writer.writerow(["Summary of Events"])
            writer.writerow(["Email Address", "Open", "Click", "Creds", "OS", "Browser"])
            # Sort campaign summary by each dict's email entry and then create results table
            target_counter = 0
            ordered_results = sorted(self.campaign_results_summary, key=lambda k: k['email'])
            for target in ordered_results:
                result = target['email']
                result += "," + target['opened']
                result += "," + target['clicked']
                result += "," + target['submitted']
                if target['email'] in self.targets_clicked:
                    for event in self.timeline:
                        if event.message == "Clicked Link" and event.email == target['email']:
                            user_agent = parse(event.details['browser']['user-agent'])
                            browser_details = user_agent.browser.family + " " + \
                                              user_agent.browser.version_string
                            os_details = user_agent.os.family + " " + user_agent.os.version_string
                            result += "," + os_details
                            result += "," + browser_details
                else:
                    result += "," + "N/A"
                    result += "," + "N/A"
                writer.writerow(["{}".format(result)])
                target_counter += 1
                print("[+] Created row for {} of {}.".format(target_counter, self.total_targets))

            print("[+] Finished writing events summary...")
            print("[+] Detailed results analysis is next and will take some time if you had a \
                  lot of targets...")
            # End of the event summary and beginning of the detailed results
            writer.writerow("")
            writer.writerow(["Detailed Analysis"])
            target_counter = 0
            for target in self.results:
                # Only create a Detailed Analysis section for targets with clicks
                if target.email in self.targets_clicked:
                    writer.writerow("")
                    writer.writerow(["{} {}".format(target.first_name, target.last_name)])
                    writer.writerow(["{}".format(target.email)])
                    # Go through all events to find events for this target
                    for event in self.timeline:
                        if event.message == "Email Sent" and event.email == target.email:
                            # Parse the timestamp into separate date and time variables
                            temp = event.time.split('T')
                            sent_date = temp[0]
                            sent_time = temp[1].split('.')[0]
                            # Record the email sent date and time in the report
                            writer.writerow(["Sent on {} at {}".format(
                                sent_date.replace(",", ""), sent_time)])

                        if event.message == "Email Opened" and event.email == target.email:
                            # Record the email preview date and time in the report
                            temp = event.time.split('T')
                            writer.writerow(["Email Preview", "{} {}".format(
                                temp[0], temp[1].split('.')[0])])

                        if event.message == "Clicked Link" and event.email == target.email:
                            temp = event.time.split('T')
                            result = temp[0] + " " + temp[1].split('.')[0]
                            # Check if browser IP matches the target's IP and record result
                            result += ",{}".format(
                                self.compare_ip_addresses(target.ip,
                                                          event.details['browser']['address'],
                                                          self.verbose))

                            # Get the location data and add to results row
                            # This is based on the IP address pulled from the browser for this event
                            # Start by getting the coordinates from GeoLite2
                            mmdb_location = self.lookup_ip(event.details['browser']['address'])
                            try:
                                if not mmdb_location is None:
                                    mmdb_latitude, mmdb_longitude = mmdb_location['location']['latitude'],\
                                                                    mmdb_location['location']['longitude']
                                    # Check if GoPhish's coordinates agree with these MMDB results
                                    result += ",{}".format(
                                        self.compare_ip_coordinates(target.latitude, target.longitude,
                                                                    mmdb_latitude, mmdb_longitude,
                                                                    event.details['browser']['address'],
                                                                    self.verbose))
                                else:
                                    result += ",IP address look-up returned None"
                            except:
                                result += ",IP address look-up returned None"

                            # Parse the user-agent string and add browser and OS details
                            user_agent = parse(event.details['browser']['user-agent'])

                            browser_details = user_agent.browser.family + " " + \
                                              user_agent.browser.version_string
                            result += ",{}".format(browser_details)
                            self.browsers.append(browser_details)

                            os_details = user_agent.os.family + " " + user_agent.os.version_string
                            result += ",{}".format(os_details)
                            self.operating_systems.append(os_details)

                            # Write the results row to the report for this target
                            writer.writerow(["Email Link Clicked"])
                            writer.writerow(["Time", "IP", "City", "Browser", "Operating System"])
                            writer.writerow([result])

                        # Now we have events for submitted data. A few notes on this:
                        # There is no expectation of data being submitted without a Clicked Link
                        # event. Assuming that, the following process does NOT flag IP
                        # mismatches or add to the list of seen locations, OSs, IPs, or browsers.
                        if event.message == "Submitted Data" and event.email == target.email:
                            temp = event.time.split('T')
                            result = temp[0] + " " + temp[1].split('.')[0]
                            result += ",{}".format(event.details['browser']['address'])

                            # Get the location data and add to results row
                            # This is based on the IP address pulled from the browser for this event
                            # Start by getting the coordinates from GeoLite2
                            mmdb_location = self.lookup_ip(event.details['browser']['address'])
                            try:
                                if not mmdb_location is None:
                                    mmdb_latitude, mmdb_longitude = mmdb_location['location']['latitude'],\
                                                                    mmdb_location['location']['longitude']
                                    # Check if GoPhish's coordinates agree with these MMDB results
                                    loc = self.compare_ip_coordinates(target.latitude,
                                                                    target.longitude, mmdb_latitude,
                                                                    mmdb_longitude,
                                                                    event.details['browser']['address'],
                                                                    self.verbose)
                                    if not loc is None:
                                        result += loc
                                    else:
                                        result += "None"
                                else:
                                    result += ",IP address look-up returned None"
                            except:
                                result += ",IP address look-up returned None"
                            # Parse the user-agent string and add browser and OS details
                            user_agent = parse(event.details['browser']['user-agent'])

                            browser_details = user_agent.browser.family + " " + \
                                              user_agent.browser.version_string
                            result += ",{}".format(browser_details)

                            os_details = user_agent.os.family + " " + user_agent.os.version_string
                            result += ",{}".format(os_details)

                            # Get just the submitted data from the event's payload
                            submitted_data = ""
                            data_payload = event.details['payload']
                            # Get all of the submitted data
                            for key, value in data_payload.items():
                                # To get just submitted data, we drop the 'rid' key
                                if not key == "rid":
                                    submitted_data += "{}:{}".format(
                                        key, str(value).strip("[").strip("]"))

                            result += ",{}".format(submitted_data)
                            # Write the results row to the report for this target
                            writer.writerow(["Submitted Data Captured"])
                            writer.writerow(["Time", "IP", "City", "Browser", "Operating System", \
                                             "Data Captured"])
                            writer.writerow([result])
                    target_counter += 1
                    print("[+] Processed detailed analysis for {} of {}.".format(
                        target_counter, self.total_targets))
                else:
                    # This target had no clicked or submitted events so move on to next
                    target_counter += 1
                    print("[+] Processed detailed analysis for {} of {}.".format(
                        target_counter, self.total_targets))
                    continue

            print("[+] Finished writing detailed analysis...")
            # End of the detailed results and the beginning of browser, location, and OS stats
            writer.writerow("")
            writer.writerow(["Recorded Browsers Based on User-Agents:"])
            writer.writerow(["Browser", "Seen"])

            counted_browsers = Counter(self.browsers)
            for key, value in counted_browsers.items():
                writer.writerow(["{},{}".format(key, value)])

            writer.writerow("")
            writer.writerow(["Record OS From Browser User-Agents:"])
            writer.writerow(["Operating System", "Seen"])

            counted_os = Counter(self.operating_systems)
            for key, value in counted_os.items():
                writer.writerow(["{},{}".format(key, value)])

            writer.writerow([" "])
            writer.writerow(["Recorded Locations from IPs:"])
            writer.writerow(["Location", "Visits"])

            counted_locations = Counter(self.locations)
            for key, value in counted_locations.items():
                writer.writerow(["{},{}".format(key, value)])

            writer.writerow([" "])
            writer.writerow(["Recorded IPs:"])
            writer.writerow(["IP Address", "Seen"])

            counted_ip_addresses = Counter(self.ip_addresses)
            for key, value in counted_ip_addresses.items():
                writer.writerow(["{},{}".format(key, value)])

            print("[+] Done! Check \'{}\' for your results.".format(self.output_csv_report))

    def write_word_report(self):
        """Assemble and output the csv file report."""
        # Create document writer using the template and a style editor
        d = Document("template.docx")
        styles = d.styles

        # Create a custom styles for table cells
        style = styles.add_style("Cell Text", WD_STYLE_TYPE.CHARACTER)
        cell_text = d.styles["Cell Text"]
        cell_text_font = cell_text.font
        cell_text_font.name = "Calibri"
        cell_text_font.size = Pt(12)
        cell_text_font.bold = True
        cell_text_font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

        style = styles.add_style("Cell Text Hit", WD_STYLE_TYPE.CHARACTER)
        cell_text_hit = d.styles["Cell Text Hit"]
        cell_text_hit_font = cell_text_hit.font
        cell_text_hit_font.name = "Calibri"
        cell_text_hit_font.size = Pt(12)
        cell_text_hit_font.bold = True
        cell_text_hit_font.color.rgb = RGBColor(0x00, 0x96, 0x00)

        style = styles.add_style("Cell Text Miss", WD_STYLE_TYPE.CHARACTER)
        cell_text_miss = d.styles["Cell Text Miss"]
        cell_text_miss_font = cell_text_miss.font
        cell_text_miss_font.name = "Calibri"
        cell_text_miss_font.size = Pt(12)
        cell_text_miss_font.bold = True
        cell_text_miss_font.color.rgb = RGBColor(0xFF, 0x00, 0x00)

        # Write a campaign summary at the top of the report
        d.add_heading("Executive Summary", 1)
        p = d.add_paragraph()
        run = p.add_run("Campaign Results For: {}".format(self.cam_name))
        run.bold = True
        # Runs are basically "runs" of text and must be aligned like we want
        # them aligned in the report -- thus they are pushed left
        if self.cam_status == "Completed":
            completed_status = self.completed_date
        else:
            completed_status = "Still Active"
        p.add_run("""
Status: {}
Created: {}
Started: {}
Completed: {}

""".format(self.cam_status, self.created_date, self.launch_date,
           completed_status))

        # Write the campaign details -- email details and template settings
        run = p.add_run("Campaign Details")
        run.bold = True
        p.add_run("""
From: {}
Subject: {}
Phish URL: {}
Redirect URL: {}
Attachment(s): {}
Captured Credentials: {}
Stored Passwords: {}

""".format(self.cam_from_address, self.cam_subject_line, self.cam_url,
           self.cam_redirect_url, self.cam_template_attachments, self.cam_capturing_credentials,
           self.cam_capturing_passwords))

        # Write a high level summary for stats
        run = p.add_run("High Level Results")
        run.bold = True
        p.add_run("""
Total Targets: {}

The following totals indicates how many events of each type GoPhish recorded:
Total Opened Events: {}
Total Clicked Events: {}
Total Submitted Data Events: {}

The following totals indicates how many targets participated in each event type:
Individuals Who Opened: {}
Individuals Who Clicked: {}
Individuals Who Submitted: {}

""".format(self.total_targets, self.total_opened, self.total_clicked,
           self.total_submitted, self.total_unique_opened, self.total_unique_clicked,
           self.total_unique_submitted))

        d.add_page_break()

        print("[+] Finished writing high level summary...")
        # End of the campaign summary and beginning of the event summary
        d.add_heading("Summary of Events", 1)
        d.add_paragraph("The following table summarizes who opened and clicked \
on emails sent in this campaign.")

        # Create a table to hold the event summary results
        table = d.add_table(rows=len(self.campaign_results_summary) + 1, cols=6, style="GoReport")

        header1 = table.cell(0, 0)
        header1.text = ""
        header1.paragraphs[0].add_run("Email Address", "Cell Text").bold = True

        header2 = table.cell(0, 1)
        header2.text = ""
        header2.paragraphs[0].add_run("Open", "Cell Text").bold = True

        header3 = table.cell(0, 2)
        header3.text = ""
        header3.paragraphs[0].add_run("Click", "Cell Text").bold = True

        header4 = table.cell(0, 3)
        header4.text = ""
        header4.paragraphs[0].add_run("Creds", "Cell Text").bold = True

        header5 = table.cell(0, 4)
        header5.text = ""
        header5.paragraphs[0].add_run("OS", "Cell Text").bold = True

        header6 = table.cell(0, 5)
        header6.text = ""
        header6.paragraphs[0].add_run("Browser", "Cell Text").bold = True

        # Sort campaign summary by each dict's email entry and then create results table
        target_counter = 0
        counter = 1
        ordered_results = sorted(self.campaign_results_summary, key=lambda k: k['email'])
        for target in ordered_results:
            email_cell = table.cell(counter, 0)
            email_cell.text = "{}".format(target['email'])

            temp_cell = table.cell(counter, 1)
            if target['opened'] == "Y":
                temp_cell.paragraphs[0].add_run(u'\u2713', "Cell Text Hit")
            else:
                temp_cell.paragraphs[0].add_run(u'\u2718', "Cell Text Miss")

            temp_cell = table.cell(counter, 2)
            if target['clicked'] == "Y":
                temp_cell.paragraphs[0].add_run(u'\u2713', "Cell Text Hit")
            else:
                temp_cell.paragraphs[0].add_run(u'\u2718', "Cell Text Miss")

            temp_cell = table.cell(counter, 3)
            if target['submitted'] == "Y":
                temp_cell.paragraphs[0].add_run(u'\u2713', "Cell Text Hit")
            else:
                temp_cell.paragraphs[0].add_run(u'\u2718', "Cell Text Miss")

            if target['email'] in self.targets_clicked:
                for event in self.timeline:
                    if event.message == "Clicked Link" and event.email == target['email']:
                        user_agent = parse(event.details['browser']['user-agent'])

                        browser_details = user_agent.browser.family + " " + \
                                          user_agent.browser.version_string
                        os_details = user_agent.os.family + " " + \
                                     user_agent.os.version_string

                        temp_cell = table.cell(counter, 4)
                        temp_cell.text = os_details
                        temp_cell = table.cell(counter, 5)
                        temp_cell.text = browser_details
            else:
                temp_cell = table.cell(counter, 4)
                temp_cell.text = "N/A"
                temp_cell = table.cell(counter, 5)
                temp_cell.text = "N/A"

            counter += 1
            target_counter += 1
            print("[+] Created table entry for {} of {}.".format(
                target_counter, self.total_targets))

        d.add_page_break()

        # End of the event summary and beginning of the detailed results
        print("[+] Finished writing events summary...")
        print("[+] Detailed results analysis is next and may take some time if you had a \
lot of targets...")
        d.add_heading("Detailed Findings", 1)
        target_counter = 0
        for target in self.results:
            # Only create a Detailed Analysis section for targets with clicks
            if target.email in self.targets_clicked:
                # Create counters to track table cell locations
                opened_counter = 1
                clicked_counter = 1
                submitted_counter = 1
                # Create section starting with a header with the first and last name
                d.add_heading("{} {}".format(target.first_name, target.last_name), 2)
                p = d.add_paragraph(target.email)
                p = d.add_paragraph()
                # Save a spot to record the email sent date and time in the report
                email_sent_run = p.add_run()
                # Go through all events to find events for this target
                for event in self.timeline:
                    if event.message == "Email Sent" and event.email == target.email:
                        # Parse the timestamp into separate date and time variables
                        # Ex: 2017-01-30T14:31:22.534880731-05:00
                        temp = event.time.split('T')
                        sent_date = temp[0]
                        sent_time = temp[1].split('.')[0]
                        # Record the email sent date and time in the run created earlier
                        email_sent_run.text = "Email sent on {} at {}".format(sent_date, sent_time)

                    if event.message == "Email Opened" and event.email == target.email:
                        if opened_counter == 1:
                            # Create the Email Opened/Previewed table
                            p = d.add_paragraph()
                            p.style = d.styles['Normal']
                            run = p.add_run("Email Previews")
                            run.bold = True

                            opened_table = d.add_table(rows=1, cols=1, style="GoReport")
                            opened_table.autofit = True
                            opened_table.allow_autofit = True

                            header1 = opened_table.cell(0, 0)
                            header1.text = ""
                            header1.paragraphs[0].add_run("Time", "Cell Text").bold = True

                        # Begin by adding a row to the table and inserting timestamp
                        opened_table.add_row()
                        timestamp = opened_table.cell(opened_counter, 0)
                        temp = event.time.split('T')
                        timestamp.text = temp[0] + " " + temp[1].split('.')[0]
                        opened_counter += 1

                    if event.message == "Clicked Link" and event.email == target.email:
                        if clicked_counter == 1:
                            # Create the Clicked Link table
                            p = d.add_paragraph()
                            p.style = d.styles['Normal']
                            run = p.add_run("Email Link Clicked")
                            run.bold = True

                            clicked_table = d.add_table(rows=1, cols=5, style="GoReport")
                            clicked_table.autofit = True
                            clicked_table.allow_autofit = True

                            header1 = clicked_table.cell(0, 0)
                            header1.text = ""
                            header1.paragraphs[0].add_run("Time", "Cell Text").bold = True

                            header2 = clicked_table.cell(0, 1)
                            header2.text = ""
                            header2.paragraphs[0].add_run("IP", "Cell Text").bold = True

                            header3 = clicked_table.cell(0, 2)
                            header3.text = ""
                            header3.paragraphs[0].add_run("City", "Cell Text").bold = True

                            header4 = clicked_table.cell(0, 3)
                            header4.text = ""
                            header4.paragraphs[0].add_run("Browser", "Cell Text").bold = True

                            header5 = clicked_table.cell(0, 4)
                            header5.text = ""
                            header5.paragraphs[0].add_run("Operating System",
                                                          "Cell Text").bold = True

                        clicked_table.add_row()
                        timestamp = clicked_table.cell(clicked_counter, 0)
                        temp = event.time.split('T')
                        timestamp.text = temp[0] + " " + temp[1].split('.')[0]

                        ip_add = clicked_table.cell(clicked_counter, 1)
                        # Check if browser IP matches the target's IP and record result
                        ip_add.text = self.compare_ip_addresses(
                            target.ip, event.details['browser']['address'], self.verbose)

                        event_location = clicked_table.cell(clicked_counter, 2)
                        # Get the location data and add to results row
                        # This is based on the IP address pulled from the browser
                        # Start by getting the coordinates from GeoLite2
                        mmdb_location = self.lookup_ip(event.details['browser']['address'])
                        try:
                            if not mmdb_location is None:
                                mmdb_latitude, mmdb_longitude = mmdb_location['location']['latitude'],\
                                                                mmdb_location['location']['longitude']
                                # Check if GoPhish's coordinates agree with these MMDB results
                                event_location.text = "{}".format(self.compare_ip_coordinates(
                                    target.latitude, target.longitude, mmdb_latitude, mmdb_longitude,
                                    event.details['browser']['address'], self.verbose))
                            else:
                                print("[!] MMDB lookup returned no location results!")
                                event_location.text = "IP address look-up returned None"
                        except:
                            print("[!] MMDB lookup returned no location results!")
                            event_location.text = "IP address look-up returned None"

                        # Parse the user-agent string for browser and OS details
                        user_agent = parse(event.details['browser']['user-agent'])

                        browser = clicked_table.cell(clicked_counter, 3)
                        browser_details = user_agent.browser.family + " " + \
                                          user_agent.browser.version_string
                        browser.text = browser_details
                        self.browsers.append(browser_details)

                        op_sys = clicked_table.cell(clicked_counter, 4)
                        os_details = user_agent.os.family + " " + user_agent.os.version_string
                        op_sys.text = os_details
                        self.operating_systems.append(os_details)

                        clicked_counter += 1

                    if event.message == "Submitted Data" and event.email == target.email:
                        if submitted_counter == 1:
                            # Create the Submitted Data table
                            p = d.add_paragraph()
                            p.style = d.styles['Normal']
                            run = p.add_run("Data Captured")
                            run.bold = True

                            submitted_table = d.add_table(rows=1, cols=6, style="GoReport")
                            submitted_table.autofit = True
                            submitted_table.allow_autofit = True

                            header1 = submitted_table.cell(0, 0)
                            header1.text = ""
                            header1.paragraphs[0].add_run("Time", "Cell Text").bold = True

                            header2 = submitted_table.cell(0, 1)
                            header2.text = ""
                            header2.paragraphs[0].add_run("IP", "Cell Text").bold = True

                            header3 = submitted_table.cell(0, 2)
                            header3.text = ""
                            header3.paragraphs[0].add_run("City", "Cell Text").bold = True

                            header4 = submitted_table.cell(0, 3)
                            header4.text = ""
                            header4.paragraphs[0].add_run("Browser", "Cell Text").bold = True

                            header5 = submitted_table.cell(0, 4)
                            header5.text = ""
                            header5.paragraphs[0].add_run("Operating System",
                                                          "Cell Text").bold = True

                            header6 = submitted_table.cell(0, 5)
                            header6.text = ""
                            header6.paragraphs[0].add_run("Data Captured",
                                                          "Cell Text").bold = True

                        submitted_table.add_row()
                        timestamp = submitted_table.cell(submitted_counter, 0)
                        temp = event.time.split('T')
                        timestamp.text = temp[0] + " " + temp[1].split('.')[0]

                        ip_add = submitted_table.cell(submitted_counter, 1)
                        ip_add.text = event.details['browser']['address']

                        event_location = submitted_table.cell(submitted_counter, 2)
                        mmdb_location = self.lookup_ip(event.details['browser']['address'])
                        try:
                            if not mmdb_location is None:
                                mmdb_latitude, mmdb_longitude = mmdb_location['location']['latitude'],\
                                                                mmdb_location['location']['longitude']
                                # Check if GoPhish's coordinates agree with these MMDB results
                                event_location.text = "{}".format(self.compare_ip_coordinates(
                                    target.latitude, target.longitude, mmdb_latitude, mmdb_longitude,
                                    event.details['browser']['address'], self.verbose))
                            else:
                                print("[!] MMDB lookup returned no location results!")
                                event_location.text = "IP address look-up returned None"
                        except:
                            print("[!] MMDB lookup returned no location results!")
                            event_location.text = "IP address look-up returned None"
                        # Parse the user-agent string and add browser and OS details
                        user_agent = parse(event.details['browser']['user-agent'])

                        browser = submitted_table.cell(submitted_counter, 3)
                        browser_details = user_agent.browser.family + " " + \
                                          user_agent.browser.version_string
                        browser.text = browser_details

                        op_sys = submitted_table.cell(submitted_counter, 4)
                        os_details = user_agent.os.family + " " + user_agent.os.version_string
                        op_sys.text = "{}".format(os_details)

                        # Get just the submitted data from the event's payload
                        submitted_data = ""
                        data = submitted_table.cell(submitted_counter, 5)
                        data_payload = event.details['payload']
                        # Get all of the submitted data
                        for key, value in data_payload.items():
                            # To get just submitted data, we drop the 'rid' key
                            if not key == "rid":
                                submitted_data += "{}:{}   ".format(
                                    key, str(value).strip("[").strip("]"))

                        data.text = "{}".format(submitted_data)

                        submitted_counter += 1
                target_counter += 1
                print("[+] Processed detailed analysis for {} of {}.".format(
                    target_counter, self.total_targets))
                d.add_page_break()
            else:
                # This target had no clicked or submitted events so move on to next
                target_counter += 1
                print("[+] Processed detailed analysis for {} of {}.".format(
                    target_counter, self.total_targets))
                continue

        print("[+] Finished writing Detailed Analysis section...")
        # End of the detailed results and the beginning of browser, location, and OS stats
        d.add_heading("Statistics", 1)
        p = d.add_paragraph("The following table shows the browsers seen:")
        # Create browser table
        browser_table = d.add_table(rows=1, cols=2, style="GoReport")
        self.set_word_column_width(browser_table.columns[0], Cm(7.24))
        self.set_word_column_width(browser_table.columns[1], Cm(3.35))

        header1 = browser_table.cell(0, 0)
        header1.text = ""
        header1.paragraphs[0].add_run("Browser", "Cell Text").bold = True

        header2 = browser_table.cell(0, 1)
        header2.text = ""
        header2.paragraphs[0].add_run("Seen", "Cell Text").bold = True

        p = d.add_paragraph("\nThe following table shows the operating systems seen:")

        # Create OS table
        os_table = d.add_table(rows=1, cols=2, style="GoReport")
        self.set_word_column_width(browser_table.columns[0], Cm(7.24))
        self.set_word_column_width(browser_table.columns[1], Cm(3.35))

        header1 = os_table.cell(0, 0)
        header1.text = ""
        header1.paragraphs[0].add_run("Operating System", "Cell Text").bold = True

        header2 = os_table.cell(0, 1)
        header2.text = ""
        header2.paragraphs[0].add_run("Seen", "Cell Text").bold = True

        p = d.add_paragraph("\nThe following table shows the locations seen:")

        # Create geo IP table
        location_table = d.add_table(rows=1, cols=2, style="GoReport")
        self.set_word_column_width(browser_table.columns[0], Cm(7.24))
        self.set_word_column_width(browser_table.columns[1], Cm(3.35))

        header1 = location_table.cell(0, 0)
        header1.text = ""
        header1.paragraphs[0].add_run("Location", "Cell Text").bold = True

        header2 = location_table.cell(0, 1)
        header2.text = ""
        header2.paragraphs[0].add_run("Visits", "Cell Text").bold = True

        p = d.add_paragraph("\nThe following table shows the IP addresses captured:")

        # Create IP address table
        ip_add_table = d.add_table(rows=1, cols=2, style="GoReport")
        self.set_word_column_width(browser_table.columns[0], Cm(7.24))
        self.set_word_column_width(browser_table.columns[1], Cm(3.35))

        header1 = ip_add_table.cell(0, 0)
        header1.text = ""
        header1.paragraphs[0].add_run("IP Address", "Cell Text").bold = True

        header2 = ip_add_table.cell(0, 1)
        header2.text = ""
        header2.paragraphs[0].add_run("Seen", "Cell Text").bold = True

        # Counters are used here again to track rows
        counter = 1
        # Counter is used to count all elements in the lists to create a unique list with totals
        counted_browsers = Counter(self.browsers)
        for key, value in counted_browsers.items():
            browser_table.add_row()
            cell = browser_table.cell(counter, 0)
            cell.text = "{}".format(key)

            cell = browser_table.cell(counter, 1)
            cell.text = "{}".format(value)
            counter += 1

        counter = 1
        counted_os = Counter(self.operating_systems)
        for key, value in counted_os.items():
            os_table.add_row()
            cell = os_table.cell(counter, 0)
            cell.text = "{}".format(key)

            cell = os_table.cell(counter, 1)
            cell.text = "{}".format(value)
            counter += 1

        counter = 1
        counted_locations = Counter(self.locations)
        for key, value in counted_locations.items():
            location_table.add_row()
            cell = location_table.cell(counter, 0)
            cell.text = "{}".format(key)

            cell = location_table.cell(counter, 1)
            cell.text = "{}".format(value)
            counter += 1

        counter = 1
        counted_ip_addresses = Counter(self.ip_addresses)
        for key, value in counted_ip_addresses.items():
            ip_add_table.add_row()
            cell = ip_add_table.cell(counter, 0)
            cell.text = "{}".format(key)

            cell = ip_add_table.cell(counter, 1)
            cell.text = "{}".format(value)
            counter += 1

        # Finalize document and save it as the value of output_word_report
        d.save("{}".format(self.output_word_report))
        print("[+] Done! Check \"{}\" for your results.".format(self.output_word_report))

    def config_section_map(self, config_parser, section):
        """This function helps by reading accepting a config file section, from gophish.config,
        and returning a dictionary object that can be referenced for configuration settings.
        """
        section_dict = {}
        options = config_parser.options(section)
        for option in options:
            try:
                section_dict[option] = config_parser.get(section, option)
                if section_dict[option] == -1:
                    print("[-] Skipping: {}".format(option))
            except:
                print("[!] There was an error with: {}".format(option))
                section_dict[option] = None
        return section_dict
